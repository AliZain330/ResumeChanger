from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, List

from docx import Document


def _is_bullet(paragraph) -> bool:
    try:
        ppr = paragraph._p.pPr
        if ppr is not None and ppr.numPr is not None:
            return True
    except Exception:
        pass
    try:
        style_name = (paragraph.style.name or "").strip().lower()
        if "bullet" in style_name:
            return True
    except Exception:
        pass
    text = (paragraph.text or "").lstrip()
    return bool(re.match(r"^(?:[-*•]\s+|\d+[\.\)]\s+)", text))


def _extract_docx_blocks(docx_path: str) -> List[Dict[str, str]]:
    document = Document(docx_path)
    blocks: List[Dict[str, str]] = []
    counter = 1

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        block_id = f"b{counter:04d}"
        counter += 1

        kind = "bullet" if _is_bullet(paragraph) else "paragraph"
        blocks.append({"id": block_id, "kind": kind, "text": text})

    return blocks


def _normalize_line(text: str) -> str:
    return " ".join(text.strip().split())


def _normalize_rewrite_text(text: str) -> str:
    # Keep content on one visual line to avoid accidental extra spacing in Word.
    return " ".join(text.replace("\r", " ").replace("\n", " ").split())


def _textual_bullet_prefix(text: str) -> str:
    match = re.match(r"^(\s*(?:[-*•]\s+|\d+[\.\)]\s+))", text or "")
    return match.group(1) if match else ""


def _strip_textual_bullet_prefix(text: str) -> str:
    return re.sub(r"^\s*(?:[-*•]\s+|\d+[\.\)]\s+)", "", text or "").strip()


def _extract_pdf_blocks(pdf_path: str) -> List[Dict[str, str]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF support requires pypdf. Install with: pip install pypdf") from exc

    reader = PdfReader(pdf_path)
    blocks: List[Dict[str, str]] = []
    counter = 1

    bullet_prefix = re.compile(r"^(?:[-*•]\s+|\d+[\.\)]\s+)")
    for page in reader.pages:
        page_text = page.extract_text() or ""
        for raw_line in page_text.splitlines():
            line = _normalize_line(raw_line)
            if not line:
                continue
            kind = "bullet" if bullet_prefix.match(line) else "paragraph"
            clean_text = bullet_prefix.sub("", line).strip() if kind == "bullet" else line
            if not clean_text:
                continue
            block_id = f"b{counter:04d}"
            counter += 1
            blocks.append({"id": block_id, "kind": kind, "text": clean_text})

    return blocks


def extract_blocks(input_path: str) -> List[Dict[str, str]]:
    suffix = Path(input_path).suffix.lower()
    if suffix == ".docx":
        return _extract_docx_blocks(input_path)
    if suffix == ".pdf":
        return _extract_pdf_blocks(input_path)
    raise ValueError(f"Unsupported file type: {suffix}")


def _replace_paragraph_text_preserving_runs(paragraph, new_text: str) -> None:
    clean_text = _normalize_rewrite_text(new_text)
    if not paragraph.runs:
        paragraph.add_run(clean_text)
        return

    # Use the first non-empty run as style anchor so font/size remain stable.
    anchor_index = 0
    for idx, run in enumerate(paragraph.runs):
        if run.text.strip():
            anchor_index = idx
            break

    paragraph.runs[anchor_index].text = clean_text
    for idx, run in enumerate(paragraph.runs):
        if idx != anchor_index:
            run.text = ""


def _apply_docx_updates(input_docx_path: str, output_docx_path: str, updates: Dict[str, str]) -> None:
    document = Document(input_docx_path)
    counter = 1

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        block_id = f"b{counter:04d}"
        counter += 1

        if block_id not in updates:
            continue

        original_text = paragraph.text or ""
        new_text = updates[block_id]

        # If user typed bullets as literal text (e.g. "• "), preserve that prefix.
        prefix = _textual_bullet_prefix(original_text)
        if prefix:
            new_text = f"{prefix}{_strip_textual_bullet_prefix(new_text)}"

        _replace_paragraph_text_preserving_runs(paragraph, new_text)

    document.save(output_docx_path)


def _apply_pdf_updates_to_docx(input_pdf_path: str, output_docx_path: str, updates: Dict[str, str]) -> None:
    blocks = _extract_pdf_blocks(input_pdf_path)
    document = Document()

    for block in blocks:
        text = updates.get(block["id"], block["text"])
        if block["kind"] == "bullet":
            document.add_paragraph(text, style="List Bullet")
        else:
            document.add_paragraph(text)

    document.save(output_docx_path)


def apply_block_updates(
    input_docx_path: str,
    output_docx_path: str,
    updates: Dict[str, str],
) -> None:
    suffix = Path(input_docx_path).suffix.lower()
    if suffix == ".docx":
        _apply_docx_updates(input_docx_path, output_docx_path, updates)
        return
    if suffix == ".pdf":
        _apply_pdf_updates_to_docx(input_docx_path, output_docx_path, updates)
        return
    raise ValueError(f"Unsupported file type: {suffix}")
